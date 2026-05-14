from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.shared import Pt


@dataclass(frozen=True)
class TemplateSection:
    title: str
    body: str
    notes: str | None = None


SIGNATURE = """Best Regards.

*Siddvin RLT*
Founder: *Mr. Vikaas D Goenka*
Call: 9117706555
www.sidvinrlt.com

Instagram: https://www.instagram.com/realtorvikasgoenka
"""


TEMPLATES: list[TemplateSection] = [
    TemplateSection(
        title="Signature (appended to every template below)",
        body=SIGNATURE,
        notes="Source: buildSignatureLines_()",
    ),
    TemplateSection(
        title="1) Property Registration (Event: PROPERTY_UPDATE)",
        body="""Dear Sir,

Property (*{propertyName}*)

The site visit to your property was done on *{visitDoneDate}*. Thank you for registering your property with us. We assure you the best services from our end. Hope to have a good relation with you. Welcome on board.

{Signature}""",
        notes="Source: buildPropertyRegistrationMessage_(payload)",
    ),
    TemplateSection(
        title="2) Brand Registration (Event: BRAND_UPDATE)",
        body="""Dear Mr. *{brandContactName}*,

Brand Name: *{brandName}*
Company Name: *{companyName}*
Designation: *{designation}*

Thank you for giving us the opportunity to join hands with you. We shall be sending you property proposals, knowledge info and regular updates through email and WhatsApp messages.

{Signature}""",
        notes="Source: buildBrandPersonMessage_(payload)",
    ),
    TemplateSection(
        title="3) Proposal (Event: PROPOSAL_UPDATED) — Property side",
        body="""Dear Sir,

We have proposed your property *{address}* to *{brandName}* of *{companyName}* on *{proposalDate}*. Next process is the site visit of *{brandName}* - *{companyName}*. For live tracking, please visit the link *{trackingLink}*.

{Signature}""",
        notes="Source: buildProposalPropertyMessage_(payload)",
    ),
    TemplateSection(
        title="4) Proposal (Event: PROPOSAL_UPDATED) — Brand side",
        body="""Dear Mr. *{brandContactName}*,

Brand Name: *{brandName}*
Company Name: *{companyName}*
Designation: *{designation}*

This is to inform you that we have proposed a commercial property *{address}* which belongs to *{ownerName}*. Please review and send us your valuable feedback.

{Signature}""",
        notes="Source: buildProposalBrandMessage_(payload)",
    ),
    TemplateSection(
        title="5) Visit Scheduled (Event: VISIT_SCHEDULED) — Property side",
        body="""Dear Sir,

Property (*{address}*)

This is to inform you that a site visit is arranged for your property on *{dateValue}* with *{brandName}* - *{companyName}*. Please keep yourself/team available for the same.

{Signature}""",
        notes="Source: buildVisitMessage_(payload, 'scheduled')",
    ),
    TemplateSection(
        title="6) Visit Scheduled (Event: VISIT_SCHEDULED) — Brand side",
        body="""Dear Mr. *{brandContactName}*,

Brand Name: *{brandName}*
Company Name: *{companyName}*

This is to inform you that the site visit of the property *{address}* which belongs to *{ownerName}* has been scheduled on *{dateValue}*.

{Signature}""",
        notes="Source: buildVisitMessage_(payload, 'scheduled')",
    ),
    TemplateSection(
        title="7) Visit Completed (Event: VISIT_COMPLETED) — Property side",
        body="""Dear Sir,

Property (*{address}*)

This is to inform you that the site visit for your property was completed by *{brandName}* - *{companyName}* on *{dateValue}*. Mr. *{ownerPresent}* from your side was present. We will keep you posted on further updates.

{Signature}""",
        notes="Source: buildVisitMessage_(payload, 'completed')",
    ),
    TemplateSection(
        title="8) Visit Completed (Event: VISIT_COMPLETED) — Brand side",
        body="""Dear Mr. *{brandContactName}*,

Brand Name: *{brandName}*
Company Name: *{companyName}*

Thank you for taking the time to visit the property *{address}* of *{ownerName}* and sharing your insights. We look forward to coordinate on the next steps.

{Signature}""",
        notes="Source: buildVisitMessage_(payload, 'completed')",
    ),
    TemplateSection(
        title="9) Follow-up: Client Meeting (Event: FOLLOW_UP_DONE) — Property side",
        body="""Dear Sir,

Property (*{address}*)

A meeting has been scheduled with Mr. *{brandPerson}* from *{brandName}* - *{companyName}* on *{dateValue}* at *{timeValue}*. Please keep yourself available.

{Signature}""",
        notes="Source: buildClientMeetingMessage_(payload)",
    ),
    TemplateSection(
        title="10) Follow-up: Client Meeting (Event: FOLLOW_UP_DONE) — Brand side",
        body="""Dear Mr. *{brandContactName}*,

Brand Name: *{brandName}*
Company Name: *{companyName}*

A meeting has been scheduled with Mr. *{ownerName}*, owner of *{address}*, on *{dateValue}* at *{timeValue}*. Please keep yourself available.

{Signature}""",
        notes="Source: buildClientMeetingMessage_(payload)",
    ),
    TemplateSection(
        title="11) Follow-up: Virtual Meeting (Event: FOLLOW_UP_DONE) — Property side",
        body="""Dear Sir,

Property (*{address}*)

A virtual meeting has been scheduled with Mr. *{brandPerson}* of *{brandName}* on *{dateValue}* at *{timeValue}*. Please make a note of the timing and join the meeting via *{linkValue}*.

{Signature}""",
        notes="Source: buildVirtualMeetingMessage_(payload)",
    ),
    TemplateSection(
        title="12) Follow-up: Virtual Meeting (Event: FOLLOW_UP_DONE) — Brand side",
        body="""Dear Mr. *{brandContactName}*,

Brand Name: *{brandName}*
Company Name: *{companyName}*

It is to inform you that a zoom meeting has been scheduled with owner *{ownerName}* of *{address}* on *{dateValue}* at *{timeValue}*. Please make a note of the timing and join the meeting via *{linkValue}*.

{Signature}""",
        notes="Source: buildVirtualMeetingMessage_(payload)",
    ),
    TemplateSection(
        title="13) Pending Terms / Agreement Meeting (Events: TERMS_AGREEMENT_UPDATED / AGREEMENT_STORE_OPENING_UPDATED) — Property side",
        body="""Dear Sir,

Property (*{address}*)

A meeting with Mr. *{brandPerson}* of *{brandName}* of *{companyName}* is scheduled on *{dateValue}* at *{timeValue}* to discuss the pending terms/agreement. Please be available for the same.

{Signature}""",
        notes="Source: buildPendingTermsAgreementMessage_(payload)",
    ),
    TemplateSection(
        title="14) Pending Terms / Agreement Meeting (Events: TERMS_AGREEMENT_UPDATED / AGREEMENT_STORE_OPENING_UPDATED) — Brand side",
        body="""Dear Mr. *{brandContactName}*,

Brand Name: *{brandName}*
Company Name: *{companyName}*

It is to inform you that the finalization of the pending terms/agreement is scheduled with *{ownerName}* of *{address}* on *{dateValue}* at *{timeValue}*. Please be available for discussion.

{Signature}""",
        notes="Source: buildPendingTermsAgreementMessage_(payload)",
    ),
    TemplateSection(
        title="15) Deal Closed / Billed (Event: SUCCESS_STORY_BILLED) — Property side",
        body="""Dear Sir,

Property (*{address}*)

We are extending our heartiest congratulations on successful completion of the lease deal with *{brandName}* - *{companyName}*. Thank you for trusting us and giving us the opportunity to work with you. Please remember us in all future endeavors.

{Signature}""",
        notes="Source: buildDealClosedMessage_(payload)",
    ),
    TemplateSection(
        title="16) Deal Closed / Billed (Event: SUCCESS_STORY_BILLED) — Brand side",
        body="""Dear Mr. *{brandContactName}*,

Brand Name: *{brandName}*
Company Name: *{companyName}*

Congratulations on successful completion of the lease deal. Your efforts and collaborations made this possible. Thank you for trusting us and giving us the opportunity to work with you. Please remember us in all future endeavors.

{Signature}""",
        notes="Source: buildDealClosedMessage_(payload)",
    ),
    TemplateSection(
        title="Appendix: Proposal Owner Compact Message (owner-only helper)",
        body="""Dear Sir,

Property: *{propertyAddress}*
Brand: *{brandName}*
Proposal Date: *{proposalDate}*
Proposal Sender: *{proposalSender}*
Brand Remarks: *{brandRemarks}*
Update By: *{updatedBy}*

{Signature}""",
        notes="Source: buildProposalOwnerCompactMessage_(payload) (helper; not wired to event router by default)",
    ),
]


def _add_mono_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)


def build_docx(output_path: Path) -> None:
    doc = Document()
    doc.add_heading("Sidvin Template Messages", level=0)
    doc.add_paragraph(
        "This document lists WhatsApp message templates found in the Sidvin Full Backend single-file script. "
        "Placeholders are shown in braces, e.g. {address}."
    )

    for section in TEMPLATES:
        doc.add_heading(section.title, level=1)
        if section.notes:
            doc.add_paragraph(section.notes)
        _add_mono_paragraph(doc, section.body)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main() -> None:
    repo_root = Path(__file__).resolve().parents[1]
    output_path = repo_root / "Sidvin_Template_Messages.docx"
    build_docx(output_path)
    print(str(output_path))


if __name__ == "__main__":
    main()

