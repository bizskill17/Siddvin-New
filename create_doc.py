import sys
import subprocess
import os

def check_and_install_docx():
    try:
        import docx
    except ImportError:
        print("Installing python-docx...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])

check_and_install_docx()

import docx

doc = docx.Document()
doc.add_heading('Sidvin Owner Notification Plan', 0)

doc.add_paragraph('This document outlines the end-to-end notification strategy for the Sidvin Owner. Every activity performed by a team member in the system will trigger a real-time notification with all non-empty fields.')

doc.add_heading('Section 1: Masters (Database Updates)', level=1)

doc.add_heading('1. Property Update', level=2)
doc.add_paragraph("Trigger: When a new Property is added or details (Rent, Area, Contact, Fee Status) are changed.", style='Intense Quote')
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run("Address: ").bold = True
p1.add_run("{address}")
p2 = doc.add_paragraph(style='List Bullet')
p2.add_run("Proposed Rent: ").bold = True
p2.add_run("{proposedRent}")
p3 = doc.add_paragraph(style='List Bullet')
p3.add_run("Proposed Area: ").bold = True
p3.add_run("{proposedArea}")
p4 = doc.add_paragraph(style='List Bullet')
p4.add_run("Floors: ").bold = True
p4.add_run("{noOfFloors} / ")
p4.add_run("Frontage: ").bold = True
p4.add_run("{frontage}")
p5 = doc.add_paragraph(style='List Bullet')
p5.add_run("Maps Link: ").bold = True
p5.add_run("{googleMapsLink}")
p6 = doc.add_paragraph(style='List Bullet')
p6.add_run("Contact: ").bold = True
p6.add_run("{contactPersonName} ({mobile})")
p7 = doc.add_paragraph(style='List Bullet')
p7.add_run("Property Fee Status: ").bold = True
p7.add_run("{propertyFeeStatus}")
p8 = doc.add_paragraph(style='List Bullet')
p8.add_run("Paper Signing Date: ").bold = True
p8.add_run("{propertyFeePaperSigningDate}")
p9 = doc.add_paragraph(style='List Bullet')
p9.add_run("Log By: ").bold = True
p9.add_run("{updatedBy}")

doc.add_heading('2. Brand Update', level=2)
doc.add_paragraph("Trigger: When a new Brand is added or details (Category, Rep, Service Fees) are changed.", style='Intense Quote')
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run("Brand Name: ").bold = True
p1.add_run("{name}")
p2 = doc.add_paragraph(style='List Bullet')
p2.add_run("Company Name: ").bold = True
p2.add_run("{companyName}")
p3 = doc.add_paragraph(style='List Bullet')
p3.add_run("Category: ").bold = True
p3.add_run("{category}")
p4 = doc.add_paragraph(style='List Bullet')
p4.add_run("Service Fee Agreed: ").bold = True
p4.add_run("{serviceFeeAgreed}")
p5 = doc.add_paragraph(style='List Bullet')
p5.add_run("Assigned Rep: ").bold = True
p5.add_run("{assignedRep}")
p6 = doc.add_paragraph(style='List Bullet')
p6.add_run("Contact: ").bold = True
p6.add_run("{contactPersonName} ({mobile})")
p7 = doc.add_paragraph(style='List Bullet')
p7.add_run("Log By: ").bold = True
p7.add_run("{updatedBy}")

doc.add_heading('3. Company / Category Master Update', level=2)
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run("List: ").bold = True
p1.add_run("{Company Master / Category Master}")
p2 = doc.add_paragraph(style='List Bullet')
p2.add_run("New Entry: ").bold = True
p2.add_run("{name}")
p3 = doc.add_paragraph(style='List Bullet')
p3.add_run("Log By: ").bold = True
p3.add_run("{updatedBy}")

doc.add_heading('4. Sidvin Team Update', level=2)
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run("Staff Name: ").bold = True
p1.add_run("{name}")
p2 = doc.add_paragraph(style='List Bullet')
p2.add_run("Designation: ").bold = True
p2.add_run("{designation}")
p3 = doc.add_paragraph(style='List Bullet')
p3.add_run("Role: ").bold = True
p3.add_run("{role}")
p4 = doc.add_paragraph(style='List Bullet')
p4.add_run("Action By: ").bold = True
p4.add_run("{Admin Name}")

doc.add_heading('Section 2: Proposals (The Deal Pipeline)', level=1)

doc.add_heading('5. Proposal Updated', level=2)
doc.add_paragraph("Trigger: Linking a Property/Brand, or editing proposal details.", style='Intense Quote')
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run("Property: ").bold = True
p1.add_run("{propertyAddress}")
p2 = doc.add_paragraph(style='List Bullet')
p2.add_run("Brand: ").bold = True
p2.add_run("{brandName}")
p3 = doc.add_paragraph(style='List Bullet')
p3.add_run("Serial No: ").bold = True
p3.add_run("{serialNo}")
p4 = doc.add_paragraph(style='List Bullet')
p4.add_run("Proposal Date: ").bold = True
p4.add_run("{proposalDate}")
p5 = doc.add_paragraph(style='List Bullet')
p5.add_run("Proposal Sender: ").bold = True
p5.add_run("{proposalSender}")
p6 = doc.add_paragraph(style='List Bullet')
p6.add_run("Brand Remarks: ").bold = True
p6.add_run("{brandRemarks}")
p7 = doc.add_paragraph(style='List Bullet')
p7.add_run("Log By: ").bold = True
p7.add_run("{updatedBy}")

doc.add_heading('6. Follow Up Done', level=2)
doc.add_paragraph("Trigger: Logging a follow-up action.", style='Intense Quote')
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run("Property: ").bold = True
p1.add_run("{propertyAddress}")
p2 = doc.add_paragraph(style='List Bullet')
p2.add_run("Brand: ").bold = True
p2.add_run("{brandName}")
p3 = doc.add_paragraph(style='List Bullet')
p3.add_run("Follow Up Date: ").bold = True
p3.add_run("{followUpDate}")
p4 = doc.add_paragraph(style='List Bullet')
p4.add_run("Status: ").bold = True
p4.add_run("{status}")
p5 = doc.add_paragraph(style='List Bullet')
p5.add_run("Remarks: ").bold = True
p5.add_run("{remarks}")
p6 = doc.add_paragraph(style='List Bullet')
p6.add_run("Next Follow Up Date: ").bold = True
p6.add_run("{nextFollowUpDate} {nextFollowUpTime}")
p7 = doc.add_paragraph(style='List Bullet')
p7.add_run("Planned Visit Date: ").bold = True
p7.add_run("{plannedVisitDate}")
p8 = doc.add_paragraph(style='List Bullet')
p8.add_run("Log By: ").bold = True
p8.add_run("{updatedBy}")

doc.add_heading('7. Proposal Cancelled', level=2)
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run("Property: ").bold = True
p1.add_run("{propertyAddress}")
p2 = doc.add_paragraph(style='List Bullet')
p2.add_run("Brand: ").bold = True
p2.add_run("{brandName}")
p3 = doc.add_paragraph(style='List Bullet')
p3.add_run("Cancellation Reason: ").bold = True
p3.add_run("{cancelRemarks}")
p4 = doc.add_paragraph(style='List Bullet')
p4.add_run("Log By: ").bold = True
p4.add_run("{updatedBy}")

doc.add_heading('8. Visit [Scheduled / Completed]', level=2)
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run("Property: ").bold = True
p1.add_run("{propertyAddress}")
p2 = doc.add_paragraph(style='List Bullet')
p2.add_run("Brand: ").bold = True
p2.add_run("{brandName}")
p3 = doc.add_paragraph(style='List Bullet')
p3.add_run("Status: ").bold = True
p3.add_run('{scheduledDate ? "Scheduled" : "Completed"}')
p4 = doc.add_paragraph(style='List Bullet')
p4.add_run("Date: ").bold = True
p4.add_run("{visitDate || scheduledDate}")
p5 = doc.add_paragraph(style='List Bullet')
p5.add_run("Visit Outcome: ").bold = True
p5.add_run("{visitOutcome}")
p6 = doc.add_paragraph(style='List Bullet')
p6.add_run("Attendance: ").bold = True
p6.add_run("{sidvinAttendees} / {brandAttendees} / {developerAttendees}")
p7 = doc.add_paragraph(style='List Bullet')
p7.add_run("Log By: ").bold = True
p7.add_run("{updatedBy}")

doc.add_heading('9. Terms / Agreement Updated', level=2)
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run("Property: ").bold = True
p1.add_run("{propertyAddress}")
p2 = doc.add_paragraph(style='List Bullet')
p2.add_run("Brand: ").bold = True
p2.add_run("{brandName}")
p3 = doc.add_paragraph(style='List Bullet')
p3.add_run("Current Stage: ").bold = True
p3.add_run('{signingDate ? "Signed" : "Finalized"}')
p4 = doc.add_paragraph(style='List Bullet')
p4.add_run("Signing Date: ").bold = True
p4.add_run("{signingDate}")
p5 = doc.add_paragraph(style='List Bullet')
p5.add_run("Finalization Date: ").bold = True
p5.add_run("{finalizationDate}")
p6 = doc.add_paragraph(style='List Bullet')
p6.add_run("Rent Commencement: ").bold = True
p6.add_run("{rentCommencementDate}")
p7 = doc.add_paragraph(style='List Bullet')
p7.add_run("Handover Date: ").bold = True
p7.add_run("{handoverDate}")
p8 = doc.add_paragraph(style='List Bullet')
p8.add_run("Agreement Registration Date: ").bold = True
p8.add_run("{agreementRegistrationDate}")
p9 = doc.add_paragraph(style='List Bullet')
p9.add_run("Planned Store Opening: ").bold = True
p9.add_run("{plannedOpeningDate} / ")
p9.add_run("Actual: ").bold = True
p9.add_run("{storeOpeningDate}")
p10 = doc.add_paragraph(style='List Bullet')
p10.add_run("Log By: ").bold = True
p10.add_run("{updatedBy}")

doc.add_heading('10. Technical Specs Finalized', level=2)
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run("Property: ").bold = True
p1.add_run("{propertyAddress}")
p2 = doc.add_paragraph(style='List Bullet')
p2.add_run("Brand: ").bold = True
p2.add_run("{brandName}")
p3 = doc.add_paragraph(style='List Bullet')
p3.add_run("AC: ").bold = True
p3.add_run("{ac} / ")
p3.add_run("Electrical Panel: ").bold = True
p3.add_run("{electricalPanel}")
p4 = doc.add_paragraph(style='List Bullet')
p4.add_run("Fire Safety: ").bold = True
p4.add_run("{fireFightingSystem} / ")
p4.add_run("Flooring: ").bold = True
p4.add_run("{flooring}")
p5 = doc.add_paragraph(style='List Bullet')
p5.add_run("Power Load: ").bold = True
p5.add_run("{powerLoad} / ")
p5.add_run("DG Backup: ").bold = True
p5.add_run("{dgBackup}")
p6 = doc.add_paragraph(style='List Bullet')
p6.add_run("Log By: ").bold = True
p6.add_run("{updatedBy}")

doc.add_heading('Section 3: Deposits (Financial Tracking)', level=1)

doc.add_heading('11. Deposit Update', level=2)
doc.add_paragraph("Trigger: When a new deposit stage is created or updated.", style='Intense Quote')
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run("Property: ").bold = True
p1.add_run("{propertyAddress}")
p2 = doc.add_paragraph(style='List Bullet')
p2.add_run("Brand: ").bold = True
p2.add_run("{brandName}")
p3 = doc.add_paragraph(style='List Bullet')
p3.add_run("Stage Name: ").bold = True
p3.add_run("{stageName}")
p4 = doc.add_paragraph(style='List Bullet')
p4.add_run("Amount Due: ").bold = True
p4.add_run("{amount}")
p5 = doc.add_paragraph(style='List Bullet')
p5.add_run("Total Received Amount: ").bold = True
p5.add_run("{receivedAmount}")
p6 = doc.add_paragraph(style='List Bullet')
p6.add_run("Latest Receipt Date: ").bold = True
p6.add_run("{receiptDate}")
p7 = doc.add_paragraph(style='List Bullet')
p7.add_run("Log By: ").bold = True
p7.add_run("{updatedBy}")

doc.add_heading('12. Receipt Update', level=2)
doc.add_paragraph("Trigger: When a new payment/receipt is added to a specific deposit stage.", style='Intense Quote')
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run("Property: ").bold = True
p1.add_run("{propertyAddress}")
p2 = doc.add_paragraph(style='List Bullet')
p2.add_run("Brand: ").bold = True
p2.add_run("{brandName}")
p3 = doc.add_paragraph(style='List Bullet')
p3.add_run("Deposit Stage: ").bold = True
p3.add_run("{stageName}")
p4 = doc.add_paragraph(style='List Bullet')
p4.add_run("Receipt Amount: ").bold = True
p4.add_run("{receiptAmount}")
p5 = doc.add_paragraph(style='List Bullet')
p5.add_run("Receipt Date: ").bold = True
p5.add_run("{receiptDate}")
p6 = doc.add_paragraph(style='List Bullet')
p6.add_run("Log By: ").bold = True
p6.add_run("{updatedBy}")

doc.add_heading('13. Success Story / Billed', level=2)
doc.add_paragraph("Trigger: Final closing and invoicing.", style='Intense Quote')
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run("Property: ").bold = True
p1.add_run("{propertyAddress}")
p2 = doc.add_paragraph(style='List Bullet')
p2.add_run("Brand: ").bold = True
p2.add_run("{brandName}")
p3 = doc.add_paragraph(style='List Bullet')
p3.add_run("Invoice Status: ").bold = True
p3.add_run("Billed")
p4 = doc.add_paragraph(style='List Bullet')
p4.add_run("Invoice No: ").bold = True
p4.add_run("{invoiceNo}")
p5 = doc.add_paragraph(style='List Bullet')
p5.add_run("Invoice Date: ").bold = True
p5.add_run("{invoiceDate}")
p6 = doc.add_paragraph(style='List Bullet')
p6.add_run("Invoice Amount: ").bold = True
p6.add_run("{invoiceAmount}")
p7 = doc.add_paragraph(style='List Bullet')
p7.add_run("Log By: ").bold = True
p7.add_run("{updatedBy}")

doc.add_heading('System Implementation Rules', level=2)
p1 = doc.add_paragraph(style='List Number')
p1.add_run("Zero-Blank Policy: Any field that is not filled by a team member in the form will not appear in the message.")
p2 = doc.add_paragraph(style='List Number')
p2.add_run("Traceability: Every message MUST capture the Property, Brand, and Team Member to ensure accountability.")
p3 = doc.add_paragraph(style='List Number')
p3.add_run("Real-Time Delivery: The notification should be triggered immediately upon form submission.")
p4 = doc.add_paragraph(style='List Number')
p4.add_run("Order: The header provides the context (e.g., Proposal Updated), while the body provides the property/brand and specific data points.")

try:
    doc.save(r'd:\Sidvin-Workflow\Sidvin_Owner_Notification_Plan.docx')
    print("Document successfully created at d:\Sidvin-Workflow\Sidvin_Owner_Notification_Plan.docx")
except Exception as e:
    print(f"Error saving document: {str(e)}")
